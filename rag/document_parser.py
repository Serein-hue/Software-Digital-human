#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档解析器 — 支持 DOCX/TXT/MD，可插拔设计。"""

import os
from typing import Any, Dict, List, Tuple

_PARSERS = {}


def _parse_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name or "标题" in style_name:
            level = 0
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            prefix = "#" * max(1, level + 1)
            paragraphs.append(f"{prefix} {text}")
        else:
            paragraphs.append(text)
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            paragraphs.append("\n".join(rows_text))
    return "\n\n".join(paragraphs)


def _parse_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _parse_md(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


_PARSERS[".docx"] = _parse_docx
_PARSERS[".txt"] = _parse_txt
_PARSERS[".md"] = _parse_md


def register_parser(extension: str, parser_fn):
    ext = extension.lower().strip()
    if not ext.startswith("."):
        ext = "." + ext
    _PARSERS[ext] = parser_fn


def parse_document(filepath: str) -> Tuple[str, Dict[str, Any]]:
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Document not found: {filepath}")
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in _PARSERS:
        raise ValueError(f"Unsupported format: {ext}. Supported: {list(_PARSERS.keys())}")
    text = _PARSERS[ext](filepath)
    metadata = {
        "source_file": os.path.basename(filepath),
        "source_path": filepath,
        "file_size": os.path.getsize(filepath),
        "file_type": ext,
        "char_count": len(text),
    }
    return text, metadata
